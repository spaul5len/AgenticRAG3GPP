from __future__ import annotations

import struct
import zlib
from pathlib import Path

from docx import Document
from docx.shared import Inches

from rag import config, figure_extractor, figure_retriever, metadata_db


def test_docx_figure_extraction_storage_and_search(tmp_path, monkeypatch):
    monkeypatch.setattr(config, "SQLITE_PATH", tmp_path / "metadata.sqlite")

    docx_path = tmp_path / "TS_33_501_sample.docx"
    image_path = tmp_path / "source.png"
    _write_png(image_path)
    _write_docx_with_figure(docx_path, image_path)

    figures = figure_extractor.extract_figures_from_docx(
        docx_path,
        figures_root=tmp_path / "figures",
        doc_type="official_spec",
        status="official",
    )

    assert len(figures) == 1
    figure = figures[0]
    assert Path(figure["image_path"]).is_file()
    assert Path(figure["image_path"]).suffix == ".png"
    assert figure["figure_number"] == "6.1.2-1"
    assert "AUSF and UDM authentication flow" in figure["caption"]
    assert "Before figure context." in figure["surrounding_text"]
    assert "After figure context." in figure["surrounding_text"]

    captured_chunks = {}

    def fake_add_chunks(collection_name, chunks, base_metadata):
        captured_chunks["collection_name"] = collection_name
        captured_chunks["chunks"] = chunks
        captured_chunks["base_metadata"] = base_metadata
        return len(chunks)

    monkeypatch.setattr(figure_extractor, "add_chunks", fake_add_chunks)

    indexed = figure_extractor.index_figure_records(figures)

    assert indexed == 1
    stored = metadata_db.list_figures(limit=10)
    assert len(stored) == 1
    assert "AUSF and UDM authentication flow" in stored[0]["caption"]
    assert captured_chunks["collection_name"] == config.FIGURE_COLLECTION
    assert "AUSF and UDM authentication flow" in captured_chunks["chunks"][0]["text"]

    def fake_search_collection(collection_name, query, k):
        assert collection_name == config.FIGURE_COLLECTION
        assert query == "AUSF UDM authentication"
        assert k == 1
        return [
            {
                "text": captured_chunks["chunks"][0]["text"],
                "metadata": captured_chunks["chunks"][0],
                "distance": 0.1,
            }
        ]

    monkeypatch.setattr(figure_retriever, "search_collection", fake_search_collection)

    results = figure_retriever.search_figures("AUSF UDM authentication", k=1)

    assert len(results) == 1
    assert results[0]["image_path"] == figure["image_path"]
    assert "AUSF and UDM authentication flow" in results[0]["caption"]


def _write_docx_with_figure(docx_path: Path, image_path: Path) -> None:
    document = Document()
    document.add_paragraph("6.1.2 Authentication architecture")
    document.add_paragraph("Before figure context.")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("Figure 6.1.2-1: AUSF and UDM authentication flow")
    document.add_paragraph("After figure context.")
    document.save(str(docx_path))


def _write_png(path: Path) -> None:
    def chunk(name: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + name
            + data
            + struct.pack(">I", zlib.crc32(name + data) & 0xFFFFFFFF)
        )

    signature = b"\x89PNG\r\n\x1a\n"
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    raw_rgb_pixel = b"\x00\x00\x6f\xc8"
    payload = (
        signature
        + chunk(b"IHDR", ihdr)
        + chunk(b"IDAT", zlib.compress(raw_rgb_pixel))
        + chunk(b"IEND", b"")
    )
    path.write_bytes(payload)
