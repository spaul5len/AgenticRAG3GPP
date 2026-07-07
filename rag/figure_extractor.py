"""Figure extraction and indexing helpers."""

from __future__ import annotations

import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document

from rag import config, metadata_db
from rag.vector_db import add_chunks


CAPTION_PATTERN = re.compile(r"^\s*(Figure|Fig\.|FIGURE)\s+(.+)", re.IGNORECASE)
FIGURE_NUMBER_PATTERN = re.compile(
    r"^\s*(?:Figure|Fig\.)\s+([0-9]+(?:[.-][0-9A-Za-z]+)*[A-Za-z]?)",
    re.IGNORECASE,
)
CLAUSE_PATTERN = re.compile(r"^\s*(\d+(?:\.\d+)+)\s+\S+")


def extract_figures_from_docx(
    path: str | Path,
    figures_root: str | Path | None = None,
    doc_type: str = "official_spec",
    status: str = "official",
) -> list[dict[str, Any]]:
    """Extract embedded DOCX images with nearby caption metadata."""

    source_file = Path(path).resolve()
    if source_file.suffix.lower() != ".docx":
        raise ValueError(f"Only DOCX figure extraction is currently supported: {source_file}")

    output_dir = _figure_output_dir(source_file, figures_root)
    output_dir.mkdir(parents=True, exist_ok=True)

    document = Document(str(source_file))
    paragraph_texts = [paragraph.text.strip() for paragraph in document.paragraphs]
    figures: list[dict[str, Any]] = []
    image_order = 0

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        for rel_id in _image_relationship_ids(paragraph):
            image_order += 1
            image_part = document.part.related_parts[rel_id]
            extension = _image_extension(image_part)
            figure_id = stable_figure_id(source_file, image_order)
            image_path = output_dir / f"{figure_id}{extension}"
            if not image_path.exists():
                image_path.write_bytes(image_part.blob)

            caption_index, caption = _nearby_caption(paragraph_texts, paragraph_index)
            figure_number = _figure_number(caption)
            surrounding_text = _surrounding_text(
                paragraph_texts, paragraph_index, caption_index
            )
            metadata = {
                "figure_id": figure_id,
                "document_title": source_file.stem,
                "document_id": metadata_db.get_document_id(source_file),
                "figure_number": figure_number,
                "caption": caption,
                "image_path": str(image_path.resolve()),
                "source_file": str(source_file),
                "page": image_order,
                "approximate_order": image_order,
                "clause": _nearest_clause(paragraph_texts, paragraph_index),
                "surrounding_text": surrounding_text,
                "doc_type": doc_type,
                "status": status,
                "created_at": _utc_now(),
            }
            figures.append(metadata)

    if figures:
        _write_manifest(output_dir, source_file, figures)
    return figures


def extract_figures_from_file(path: str | Path, **kwargs: Any) -> list[dict[str, Any]]:
    """Dispatch figure extraction by file type."""

    file_path = Path(path)
    if file_path.suffix.lower() == ".docx":
        return extract_figures_from_docx(file_path, **kwargs)
    return []


def index_figure_records(figures: list[dict[str, Any]]) -> int:
    """Register and index extracted figures into SQLite and Chroma."""

    chunks: list[dict[str, Any]] = []
    for figure in figures:
        metadata_db.register_figure(figure)
        chunks.append(
            {
                "chunk_id": figure["figure_id"],
                "doc_id": figure.get("source_file") or figure.get("document_title"),
                "text": figure_search_text(figure),
                "page": figure.get("page") or figure.get("approximate_order"),
                "figure_id": figure["figure_id"],
                "figure_number": figure.get("figure_number"),
                "caption": figure.get("caption"),
                "image_path": figure.get("image_path"),
                "source_file": figure.get("source_file"),
                "clause": figure.get("clause"),
                "doc_type": figure.get("doc_type"),
                "status": figure.get("status"),
                "approximate_order": figure.get("approximate_order"),
            }
        )

    if not chunks:
        return 0

    base_metadata = {
        "doc_type": "figure",
        "status": "figure_evidence",
        "collection_name": config.FIGURE_COLLECTION,
        "doc_id": "figures",
    }
    return add_chunks(config.FIGURE_COLLECTION, chunks, base_metadata)


def figure_search_text(figure: dict[str, Any]) -> str:
    """Build searchable text for a figure."""

    parts = [
        figure.get("document_title"),
        figure.get("figure_number"),
        figure.get("caption"),
        figure.get("clause"),
        figure.get("surrounding_text"),
    ]
    return "\n".join(str(part) for part in parts if part)


def stable_figure_id(source_file: Path, order: int) -> str:
    stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", source_file.stem).strip("_") or "document"
    return f"{stem}-fig-{order:04d}"


def _image_relationship_ids(paragraph) -> list[str]:
    rel_ids: list[str] = []
    for blip in paragraph._element.xpath(".//*[local-name()='blip']"):
        rel_id = blip.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )
        if rel_id:
            rel_ids.append(rel_id)
    return rel_ids


def _image_extension(image_part) -> str:
    suffix = Path(str(image_part.partname)).suffix
    if suffix:
        return suffix.lower()
    content_type = str(getattr(image_part, "content_type", "")).lower()
    if "jpeg" in content_type:
        return ".jpg"
    if "png" in content_type:
        return ".png"
    if "gif" in content_type:
        return ".gif"
    return ".bin"


def _nearby_caption(
    paragraphs: list[str], paragraph_index: int
) -> tuple[int | None, str]:
    for offset in (1, -1, 2, -2):
        index = paragraph_index + offset
        if 0 <= index < len(paragraphs) and CAPTION_PATTERN.match(paragraphs[index]):
            return index, paragraphs[index]
    return None, ""


def _figure_number(caption: str) -> str | None:
    match = FIGURE_NUMBER_PATTERN.match(caption or "")
    return match.group(1) if match else None


def _surrounding_text(
    paragraphs: list[str], paragraph_index: int, caption_index: int | None
) -> str:
    selected: list[str] = []
    for index in (
        paragraph_index - 1,
        caption_index,
        paragraph_index + 1,
        None if caption_index is None else caption_index - 1,
        None if caption_index is None else caption_index + 1,
    ):
        if index is None or not 0 <= index < len(paragraphs):
            continue
        text = paragraphs[index].strip()
        if text and text not in selected:
            selected.append(text)
    return "\n".join(selected)


def _nearest_clause(paragraphs: list[str], paragraph_index: int) -> str | None:
    for index in range(paragraph_index, -1, -1):
        match = CLAUSE_PATTERN.match(paragraphs[index])
        if match:
            return match.group(1)
    return None


def _figure_output_dir(source_file: Path, figures_root: str | Path | None) -> Path:
    root = Path(figures_root) if figures_root else config.DATA_DIR / "figures"
    return root / source_file.stem


def _write_manifest(
    output_dir: Path, source_file: Path, figures: list[dict[str, Any]]
) -> None:
    manifest_path = output_dir / "figures.json"
    payload = {
        "source_file": str(source_file),
        "figures": figures,
    }
    manifest_path.write_text(json.dumps(payload, indent=2, sort_keys=True), encoding="utf-8")


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")
